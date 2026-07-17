from typing import Any
import os
import tempfile
from html import escape

import fitz
from docx import Document

from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse, StreamingResponse
from openai import OpenAI
from pydantic import BaseModel, Field
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.platypus import Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from io import BytesIO


LOCAL_LLM_BASE_URL = "http://127.0.0.1:11434/v1"
MODEL_ID = "llama3"

client = OpenAI(
    base_url=LOCAL_LLM_BASE_URL,
    api_key="ollama",
)

app = FastAPI(
    title="Legal Clause Analyzer",
    version="1.2",
    description=(
        "Privacy-first legal clause analyzer with local LLM support. "
        "Supports PDF, DOCX and Contract Comparison reports."
    ),
)
latest_analysis = {}
latest_comparison = {}


class AnalyzeRequest(BaseModel):
    contract_text: str = Field(
        ...,
        min_length=20,
        description="Contract text to analyze.",
    )
    use_llm: bool = Field(
        default=False,
        description="Use local Ollama LLM for a short summary.",
    )


class Finding(BaseModel):
    clause_type: str
    risk_level: str
    matched_keywords: list[str]
    explanation: str


class RiskScores(BaseModel):
    overall_risk_score: int
    gdpr_readiness_score: int
    eu_ai_act_readiness_score: int


class AIActCheck(BaseModel):
    ai_act_triggered: bool
    matched_ai_terms: list[str]
    matched_high_risk_terms: list[str]
    missing_controls: list[str]
    issues: list[str]
    recommendations: list[str]
    summary: str


class GDPRCheck(BaseModel):
    personal_data_detected: bool
    sensitive_data_detected: bool
    matched_personal_data_terms: list[str]
    matched_sensitive_data_terms: list[str]
    missing_controls: list[str]
    issues: list[str]
    recommendations: list[str]


class AnalysisResult(BaseModel):
    findings: list[Finding]
    risk_scores: RiskScores
    ai_act_check: AIActCheck
    gdpr_check: GDPRCheck
    llm_summary: str | None


class ClauseComparison(BaseModel):
    contract_a_clause_types: list[str]
    contract_b_clause_types: list[str]
    added: list[str]
    removed: list[str]
    common: list[str]


class ScoreDelta(BaseModel):
    contract_a: int
    contract_b: int
    difference: int
    trend: str


class ScoreComparison(BaseModel):
    overall_risk_score: ScoreDelta
    gdpr_readiness_score: ScoreDelta
    eu_ai_act_readiness_score: ScoreDelta


class ComparisonSummary(BaseModel):
    added_clauses_count: int
    removed_clauses_count: int
    common_clauses_count: int
    total_clauses_contract_a: int
    total_clauses_contract_b: int


class Comparison(BaseModel):
    clause_comparison: ClauseComparison
    score_comparison: ScoreComparison
    summary: ComparisonSummary


class CharactersExtracted(BaseModel):
    contract_a: int
    contract_b: int
    total: int


class CompareContractsResponse(BaseModel):
    project: str
    contract_a_filename: str
    contract_b_filename: str
    characters_extracted: CharactersExtracted
    contract_a_analysis: AnalysisResult
    contract_b_analysis: AnalysisResult
    comparison: Comparison
    disclaimer: str


CLAUSE_RULES: list[dict[str, Any]] = [
    {
        "name": "Force Majeure",
        "keywords": [
            "force majeure",
            "act of god",
            "unforeseeable event",
            "natural disaster",
        ],
        "risk_level": "Medium",
        "explanation": (
            "Review notice duties, suspension rights, and termination rights."
        ),
    },
    {
        "name": "Liability Limitation",
        "keywords": [
            "liability",
            "limit liability",
            "limitation of liability",
            "cap on liability",
        ],
        "risk_level": "High",
        "explanation": (
            "Liability may be limited. Check whether remedies are one-sided."
        ),
    },
    {
        "name": "Termination",
        "keywords": [
            "termination",
            "terminate this agreement",
            "termination for convenience",
            "termination for cause",
        ],
        "risk_level": "Medium",
        "explanation": (
            "Check notice periods, cure periods, and post-termination duties."
        ),
    },
    {
        "name": "Confidentiality",
        "keywords": [
            "confidential information",
            "confidentiality",
            "non-disclosure",
            "trade secret",
        ],
        "risk_level": "Medium",
        "explanation": (
            "Check scope, duration, exceptions, and return or deletion duties."
        ),
    },
    {
        "name": "Data Protection",
        "keywords": [
            "personal data",
            "data protection",
            "gdpr",
            "data subject",
            "processor",
            "controller",
        ],
        "risk_level": "High",
        "explanation": (
            "Personal data may be involved. Check GDPR roles and safeguards."
        ),
    },
    {
        "name": "AI Systems",
        "keywords": [
            "ai system",
            "artificial intelligence",
            "machine learning",
            "automated decision",
            "automated decision-making",
            "algorithmic decision",
        ],
        "risk_level": "High",
        "explanation": (
            "AI usage detected. Review EU AI Act and data governance duties."
        ),
    },
]

AI_SYSTEM_KEYWORDS: list[str] = [
    "ai system",
    "artificial intelligence",
    "machine learning",
    "automated decision",
    "automated decision-making",
    "algorithmic decision",
    "predictive model",
]

HIGH_RISK_AI_KEYWORDS: list[str] = [
    "employment",
    "worker management",
    "credit scoring",
    "education",
    "biometric identification",
    "law enforcement",
    "migration",
    "essential private services",
    "essential public services",
]

PERSONAL_DATA_KEYWORDS: list[str] = [
    "personal data",
    "personally identifiable information",
    "pii",
    "data subject",
    "customer data",
    "employee data",
    "user data",
]

SENSITIVE_DATA_KEYWORDS: list[str] = [
    "health data",
    "biometric data",
    "genetic data",
    "racial or ethnic origin",
    "political opinions",
    "religious beliefs",
    "trade union",
    "sexual orientation",
]

AI_ACT_CONTROL_CHECKS: dict[str, list[str]] = {
    "human oversight": [
        "human oversight",
        "human review",
        "manual review",
        "human-in-the-loop",
    ],
    "transparency": [
        "transparency",
        "user notice",
        "disclosure",
        "explainability",
    ],
    "logging": [
        "logging",
        "logs",
        "audit trail",
        "record keeping",
    ],
    "risk management": [
        "risk management",
        "risk assessment",
        "mitigation measures",
    ],
    "data governance": [
        "data governance",
        "training data",
        "data quality",
        "dataset",
    ],
    "incident reporting": [
        "incident reporting",
        "serious incident",
        "notification duty",
    ],
}

GDPR_CONTROL_CHECKS: dict[str, list[str]] = {
    "lawful basis": [
        "lawful basis",
        "consent",
        "legitimate interest",
        "contractual necessity",
    ],
    "retention": [
        "retention",
        "deletion",
        "erase",
        "storage period",
    ],
    "security": [
        "encryption",
        "access control",
        "security measures",
        "confidentiality",
    ],
    "data subject rights": [
        "data subject rights",
        "access request",
        "rectification",
        "erasure",
        "objection",
    ],
    "processor/controller roles": [
        "controller",
        "processor",
        "data processing agreement",
        "dpa",
    ],
}


def match_keywords(text: str, keywords: list[str]) -> list[str]:
    return [keyword for keyword in keywords if keyword in text]


def has_any(text: str, keywords: list[str]) -> bool:
    return any(keyword in text for keyword in keywords)


def find_missing_controls(
    lower_text: str,
    controls: dict[str, list[str]],
) -> list[str]:
    missing: list[str] = []

    for control_name, keywords in controls.items():
        if not has_any(lower_text, keywords):
            missing.append(control_name)

    return missing


def detect_clauses(text: str) -> list[dict[str, Any]]:
    lower_text = text.lower()
    findings: list[dict[str, Any]] = []

    for rule in CLAUSE_RULES:
        matched = match_keywords(lower_text, rule["keywords"])

        if matched:
            findings.append(
                {
                    "clause_type": rule["name"],
                    "risk_level": rule["risk_level"],
                    "matched_keywords": matched,
                    "explanation": rule["explanation"],
                }
            )

    return findings


def ai_act_compliance_check(text: str) -> dict[str, Any]:
    lower_text = text.lower()

    ai_terms = match_keywords(lower_text, AI_SYSTEM_KEYWORDS)
    high_risk_terms = match_keywords(lower_text, HIGH_RISK_AI_KEYWORDS)
    missing_controls = find_missing_controls(
        lower_text,
        AI_ACT_CONTROL_CHECKS,
    )

    ai_act_triggered = bool(ai_terms)
    issues: list[str] = []
    recommendations: list[str] = []

    if ai_act_triggered:
        summary = (
            "AI-related language was detected. This is a readiness check, "
            "not a final legal compliance determination."
        )

        if high_risk_terms:
            issues.append(
                "Possible high-risk AI context detected. "
                "Human legal review is recommended."
            )

        if missing_controls:
            issues.append(
                "The contract may not clearly address: "
                + ", ".join(missing_controls)
                + "."
            )

        recommendations.extend(
            [
                "Clarify provider and deployer responsibilities.",
                "Add human oversight language where relevant.",
                "Document logging and audit-trail obligations.",
                "Review data governance and training-data clauses.",
                "Add transparency and user-notice obligations if needed.",
            ]
        )
    else:
        summary = (
            "No obvious AI-system language was detected. "
            "AI Act obligations are not clearly triggered by this text."
        )
        recommendations.append(
            "Keep this as a compliance-readiness check, not legal advice."
        )

    if not issues:
        issues.append(
            "No obvious AI Act issue was detected at keyword level."
        )

    return {
        "ai_act_triggered": ai_act_triggered,
        "matched_ai_terms": ai_terms,
        "matched_high_risk_terms": high_risk_terms,
        "missing_controls": missing_controls if ai_act_triggered else [],
        "issues": issues,
        "recommendations": recommendations,
        "summary": summary,
    }


def gdpr_privacy_check(text: str) -> dict[str, Any]:
    lower_text = text.lower()

    personal_data_terms = match_keywords(
        lower_text,
        PERSONAL_DATA_KEYWORDS,
    )
    sensitive_data_terms = match_keywords(
        lower_text,
        SENSITIVE_DATA_KEYWORDS,
    )
    missing_controls = find_missing_controls(
        lower_text,
        GDPR_CONTROL_CHECKS,
    )

    personal_data_detected = bool(personal_data_terms)
    sensitive_data_detected = bool(sensitive_data_terms)

    issues: list[str] = []
    recommendations: list[str] = []

    if personal_data_detected:
        issues.append(
            "Personal data language detected. GDPR safeguards should be "
            "reviewed."
        )

        if missing_controls:
            issues.append(
                "The contract may not clearly address: "
                + ", ".join(missing_controls)
                + "."
            )

        recommendations.extend(
            [
                "Clarify controller and processor roles.",
                "Add retention and deletion obligations.",
                "Include access control and encryption language.",
                "Address data subject rights where relevant.",
                "Avoid using real personal data in public demos.",
            ]
        )
    else:
        recommendations.append(
            "No obvious personal-data language detected at keyword level."
        )

    if sensitive_data_detected:
        issues.append(
            "Sensitive personal data may be involved. "
            "Apply stricter review and minimization controls."
        )

    if not issues:
        issues.append(
            "No obvious GDPR issue was detected at keyword level."
        )

    return {
        "personal_data_detected": personal_data_detected,
        "sensitive_data_detected": sensitive_data_detected,
        "matched_personal_data_terms": personal_data_terms,
        "matched_sensitive_data_terms": sensitive_data_terms,
        "missing_controls": missing_controls if personal_data_detected else [],
        "issues": issues,
        "recommendations": recommendations,
    }


def generate_llm_summary(
    text: str,
    findings: list[dict[str, Any]],
    ai_act_check: dict[str, Any],
    gdpr_check: dict[str, Any],
) -> str:
    prompt = f"""
You are an AI LegalTech Assistant specialized in:

- GDPR compliance
- EU AI Act compliance
- AI Governance
- Commercial contract review

This is NOT legal advice.

Based ONLY on the supplied analysis, generate a professional compliance report.

Use EXACTLY the following structure:

# Executive Summary

# Overall Risk Level

# GDPR Readiness

# EU AI Act Readiness

# Key Legal Risks

# Missing Compliance Controls

# Priority Recommendations
(List only the five most important recommendations.)

# Conclusion

# Disclaimer
State clearly that this is a compliance-readiness assessment and not legal advice.

Do not repeat the contract.

Do not invent facts.

Use only the supplied findings.

-------------------------

Clause Findings

{findings}

-------------------------

EU AI Act Analysis

{ai_act_check}

-------------------------

GDPR Analysis

{gdpr_check}

-------------------------

Contract Excerpt

{text[:6000]}
"""

    try:
        completion = client.chat.completions.create(
            model=MODEL_ID,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a cautious legal technology assistant. "
                        "You explain risks clearly and avoid legal advice."
                    ),
                },
                {
                    "role": "user",
                    "content": prompt,
                },
            ],
            temperature=0.1,
            top_p=0.9,
            max_tokens=700,
        )

        content = completion.choices[0].message.content
        return content or "Local LLM returned an empty response."

    except Exception as exc:
        return (
            "Local LLM summary unavailable. "
            f"Reason: {exc.__class__.__name__}: {exc}"
        )


def calculate_risk_score(
    findings: list[dict[str, Any]],
    ai_act_check: dict[str, Any],
    gdpr_check: dict[str, Any],
) -> dict[str, int]:

    score = 0

    for finding in findings:
        if finding["risk_level"] == "High":
            score += 15
        elif finding["risk_level"] == "Medium":
            score += 8
        else:
            score += 3

    score += len(ai_act_check["missing_controls"]) * 4
    score += len(gdpr_check["missing_controls"]) * 4

    score = min(score, 100)

    return {
        "overall_risk_score": score,
        "gdpr_readiness_score": max(
            0,
            100 - len(gdpr_check["missing_controls"]) * 10,
        ),
        "eu_ai_act_readiness_score": max(
            0,
            100 - len(ai_act_check["missing_controls"]) * 10,
        ),
    }


def run_full_analysis(
    contract_text: str,
    use_llm: bool,
) -> dict[str, Any]:
    findings = detect_clauses(contract_text)

    ai_act_check = ai_act_compliance_check(contract_text)

    gdpr_check = gdpr_privacy_check(contract_text)

    risk_scores = calculate_risk_score(
        findings,
        ai_act_check,
        gdpr_check,
    )

    llm_summary = None

    if use_llm:
        llm_summary = generate_llm_summary(
            contract_text,
            findings,
            ai_act_check,
            gdpr_check,
        )

    return {
        "findings": findings,
        "risk_scores": risk_scores,
        "ai_act_check": ai_act_check,
        "gdpr_check": gdpr_check,
        "llm_summary": llm_summary,
    }


def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    """Extract text from PDF bytes."""

    document = fitz.open(stream=pdf_bytes, filetype="pdf")

    text = ""

    for page in document:
        text += page.get_text()

    document.close()

    return text


def extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    """Extract text from DOCX bytes, including paragraphs and tables."""

    document = Document(BytesIO(docx_bytes))

    text_parts: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            if any(row_text):
                text_parts.append(" | ".join(row_text))

    return "\n".join(text_parts)


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Dispatch to the correct extractor based on file extension."""

    lower_name = filename.lower()

    if lower_name.endswith(".pdf"):
        return extract_text_from_pdf_bytes(file_bytes)

    if lower_name.endswith(".docx"):
        return extract_text_from_docx_bytes(file_bytes)

    raise ValueError(
        f"Unsupported file format: {filename}. "
        "Only PDF and DOCX files are currently supported."
    )


@app.post("/analyze-pdf")
async def analyze_pdf(
    file: UploadFile = File(...),
    use_llm: bool = False,
) -> dict[str, Any]:

    global latest_analysis 

    pdf_bytes = await file.read()

    contract_text = extract_text(pdf_bytes, file.filename)

    result = run_full_analysis(contract_text, use_llm)

    latest_analysis = {
        "findings": result["findings"],
        "risk_scores": result["risk_scores"],
        "ai_act_check": result["ai_act_check"],
        "gdpr_check": result["gdpr_check"],
        "llm_summary": result["llm_summary"],
    }


    return {
        "project": "Legal Clause Analyzer",
        "source": file.filename,
        "characters_analyzed": len(contract_text),
        "clause_findings": result["findings"],
        "ai_act_compliance_check": result["ai_act_check"],
        "gdpr_privacy_check": result["gdpr_check"],
        "risk_scores": result["risk_scores"],
        "llm_summary": result["llm_summary"],
        "disclaimer": (
            "This output is for compliance-readiness and "
            "demonstration only. It is not legal advice."
        ),
    }


@app.post("/analyze-docx")
async def analyze_docx(
    file: UploadFile = File(...),
    use_llm: bool = False,
) -> dict[str, Any]:

    global latest_analysis

    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(
            status_code=400,
            detail="Only DOCX files are supported.",
        )

    docx_bytes = await file.read()

    contract_text = extract_text(docx_bytes, file.filename)

    result = run_full_analysis(contract_text, use_llm)

    latest_analysis = {
        "findings": result["findings"],
        "risk_scores": result["risk_scores"],
        "ai_act_check": result["ai_act_check"],
        "gdpr_check": result["gdpr_check"],
        "llm_summary": result["llm_summary"],
    }

    return {
        "project": "Legal Clause Analyzer",
        "source": file.filename,
        "characters_analyzed": len(contract_text),
        "clause_findings": result["findings"],
        "ai_act_compliance_check": result["ai_act_check"],
        "gdpr_privacy_check": result["gdpr_check"],
        "risk_scores": result["risk_scores"],
        "llm_summary": result["llm_summary"],
        "disclaimer": (
            "This output is for compliance-readiness and "
            "demonstration only. It is not legal advice."
        ),
    }


def compare_analysis_results(
    result_a: dict[str, Any],
    result_b: dict[str, Any],
) -> dict[str, Any]:
    """Compare two full analysis results.

    Returns a structured comparison of detected clauses and risk/readiness
    scores without generating PDFs or LLM summaries.
    """

    findings_a = result_a.get("findings", [])
    findings_b = result_b.get("findings", [])

    clauses_a = {f["clause_type"] for f in findings_a}
    clauses_b = {f["clause_type"] for f in findings_b}

    added = sorted(clauses_b - clauses_a)
    removed = sorted(clauses_a - clauses_b)
    common = sorted(clauses_a & clauses_b)

    scores_a = result_a.get("risk_scores", {})
    scores_b = result_b.get("risk_scores", {})

    def score_delta(key: str) -> dict[str, Any]:
        value_a = scores_a.get(key, 0)
        value_b = scores_b.get(key, 0)
        diff = value_b - value_a

        if diff > 0:
            trend = "increased"
        elif diff < 0:
            trend = "decreased"
        else:
            trend = "unchanged"

        return {
            "contract_a": value_a,
            "contract_b": value_b,
            "difference": diff,
            "trend": trend,
        }

    return {
        "clause_comparison": {
            "contract_a_clause_types": sorted(clauses_a),
            "contract_b_clause_types": sorted(clauses_b),
            "added": added,
            "removed": removed,
            "common": common,
        },
        "score_comparison": {
            "overall_risk_score": score_delta("overall_risk_score"),
            "gdpr_readiness_score": score_delta("gdpr_readiness_score"),
            "eu_ai_act_readiness_score": score_delta(
                "eu_ai_act_readiness_score"
            ),
        },
        "summary": {
            "added_clauses_count": len(added),
            "removed_clauses_count": len(removed),
            "common_clauses_count": len(common),
            "total_clauses_contract_a": len(findings_a),
            "total_clauses_contract_b": len(findings_b),
        },
    }


def _join_items(items: list[Any], fallback: str = "—") -> str:
    """Return a comma-separated string or a fallback placeholder."""
    if not items:
        return fallback
    return ", ".join(str(item) for item in items)


def _risk_score_color(score: int) -> Any:
    """Color for an overall risk score (higher is worse)."""
    if score >= 70:
        return colors.red
    if score >= 40:
        return colors.orange
    return colors.green


def _readiness_score_color(score: int) -> Any:
    """Color for a readiness score (higher is better)."""
    if score >= 80:
        return colors.green
    if score >= 50:
        return colors.orange
    return colors.red


def _paragraph_cell(
    text: Any,
    style: Any,
    bold: bool = False,
    text_color: Any = None,
) -> Paragraph:
    """Return a ReportLab Paragraph with safe HTML escaping."""
    if text is None:
        text = "—"
    safe_text = str(text)
    safe_text = escape(safe_text)
    safe_text = safe_text.replace("\n", "<br/>")
    if bold:
        safe_text = f"<b>{safe_text}</b>"
    if text_color is not None:
        style = style.clone("_colored", textColor=text_color)
    return Paragraph(safe_text, style)


def _pdf_heading(
    story: list[Any],
    styles: Any,
    text: str,
    style: str = "Heading2",
) -> None:
    story.append(Paragraph(text, styles[style]))


def _pdf_body(
    story: list[Any],
    styles: Any,
    text: str,
) -> None:
    story.append(Paragraph(text, styles["BodyText"]))


def _pdf_bullets(
    story: list[Any],
    styles: Any,
    items: list[str],
    prefix: str = "• ",
) -> None:
    for item in items:
        story.append(Paragraph(f"{prefix}{item}", styles["BodyText"]))


def _pdf_labeled_bullets(
    story: list[Any],
    styles: Any,
    title: str,
    sections: list[tuple[str, list[str]]],
    fallback: str = "• None detected.",
) -> None:
    _pdf_heading(story, styles, title, "Heading3")
    for label, items in sections:
        _pdf_body(story, styles, f"<b>Contract {label}</b>")
        if items:
            _pdf_bullets(story, styles, items)
        else:
            _pdf_body(story, styles, fallback)


def _base_table_style(align: str = "LEFT") -> TableStyle:
    """Return the common table style used across PDF reports."""
    return TableStyle(
        [
            ("BACKGROUND", (0, 0), (-1, 0), colors.darkblue),
            ("TEXTCOLOR", (0, 0), (-1, -1), colors.black),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("BACKGROUND", (0, 1), (-1, -1), colors.beige),
            ("GRID", (0, 0), (-1, -1), 1, colors.grey),
            ("BOX", (0, 0), (-1, -1), 1, colors.black),
            ("ALIGN", (0, 0), (-1, -1), align),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]
    )


def _make_table(
    rows: list[list[Any]],
    colWidths: list[Any],
    styles: Any,
    align: str = "LEFT",
    first_col_bold: bool = True,
    extra_commands: list[tuple[Any, ...]] | None = None,
    repeat_rows: int = 1,
) -> Table:
    """Build a styled ReportLab Table with Paragraph-based wrapping cells."""
    alignment_map = {
        "LEFT": TA_LEFT,
        "CENTER": TA_CENTER,
        "RIGHT": TA_RIGHT,
    }
    body_alignment = alignment_map.get(align, TA_LEFT)
    body_style = styles["BodyText"].clone(
        "_table_body",
        alignment=body_alignment,
    )
    header_style = body_style.clone(
        "_table_header",
        textColor=colors.white,
        fontName="Helvetica-Bold",
    )

    wrapped_rows: list[list[Paragraph]] = []

    for row_idx, row in enumerate(rows):
        wrapped_row: list[Paragraph] = []
        for col_idx, value in enumerate(row):
            if isinstance(value, Paragraph):
                wrapped_row.append(value)
            elif row_idx == 0:
                wrapped_row.append(_paragraph_cell(value, header_style))
            else:
                is_bold = first_col_bold and col_idx == 0
                wrapped_row.append(
                    _paragraph_cell(value, body_style, bold=is_bold)
                )
        wrapped_rows.append(wrapped_row)

    style = _base_table_style(align)
    for command in extra_commands or []:
        style.add(*command)
    table = Table(wrapped_rows, colWidths=colWidths, repeatRows=repeat_rows)
    table.setStyle(style)
    return table


def _contract_summary_table(
    filename: str,
    findings: list[dict[str, Any]],
    risk_scores: dict[str, int],
    gdpr_check: dict[str, Any],
    ai_act_check: dict[str, Any],
    styles: Any,
    usable_width: float,
) -> Table:
    """Build the attribute/value table for a single contract summary."""
    label_width = 1.5 * inch
    value_width = usable_width - label_width
    body_style = styles["BodyText"]

    overall_score = risk_scores.get("overall_risk_score", 0)
    gdpr_score = risk_scores.get("gdpr_readiness_score", 0)
    ai_score = risk_scores.get("eu_ai_act_readiness_score", 0)

    rows: list[list[Any]] = [
        ["Attribute", "Value"],
        ["Filename", filename],
        ["Detected clauses", str(len(findings))],
        ["Clause types", _join_items([f["clause_type"] for f in findings])],
        [
            "Overall risk score",
            _paragraph_cell(
                str(overall_score),
                body_style,
                text_color=_risk_score_color(overall_score),
            ),
        ],
        [
            "GDPR readiness score",
            _paragraph_cell(
                str(gdpr_score),
                body_style,
                text_color=_readiness_score_color(gdpr_score),
            ),
        ],
        [
            "EU AI Act readiness score",
            _paragraph_cell(
                str(ai_score),
                body_style,
                text_color=_readiness_score_color(ai_score),
            ),
        ],
        [
            "Personal data detected",
            "Yes" if gdpr_check.get("personal_data_detected") else "No",
        ],
        [
            "Sensitive data detected",
            "Yes" if gdpr_check.get("sensitive_data_detected") else "No",
        ],
        [
            "GDPR missing controls",
            _join_items(gdpr_check.get("missing_controls", [])),
        ],
        [
            "AI Act triggered",
            "Yes" if ai_act_check.get("ai_act_triggered") else "No",
        ],
        [
            "High-risk AI terms",
            _join_items(ai_act_check.get("matched_high_risk_terms", [])),
        ],
        [
            "AI Act missing controls",
            _join_items(ai_act_check.get("missing_controls", [])),
        ],
    ]
    return _make_table(
        rows,
        colWidths=[label_width, value_width],
        styles=styles,
    )


def _score_comparison_table(
    score_comparison: dict[str, Any],
    styles: Any,
    usable_width: float,
) -> Table:
    """Build the side-by-side risk/readiness score table with colors."""
    label_width = 1.5 * inch
    value_width = (usable_width - label_width) / 4

    body_style_left = styles["BodyText"].clone(
        "_score_label",
        alignment=TA_LEFT,
    )
    body_style_center = styles["BodyText"].clone(
        "_score_value",
        alignment=TA_CENTER,
    )

    def _row(metric_key: str, label: str) -> list[Any]:
        delta = score_comparison.get(metric_key, {})
        value_a = delta.get("contract_a", 0)
        value_b = delta.get("contract_b", 0)
        diff = delta.get("difference", 0)
        trend = delta.get("trend", "unchanged")

        if metric_key == "overall_risk_score":
            color_a = _risk_score_color(value_a)
            color_b = _risk_score_color(value_b)
            color_diff = colors.red if diff > 0 else colors.green
        else:
            color_a = _readiness_score_color(value_a)
            color_b = _readiness_score_color(value_b)
            color_diff = colors.green if diff > 0 else colors.red

        return [
            _paragraph_cell(label, body_style_left, bold=True),
            _paragraph_cell(
                str(value_a), body_style_center, text_color=color_a
            ),
            _paragraph_cell(
                str(value_b), body_style_center, text_color=color_b
            ),
            _paragraph_cell(
                str(diff), body_style_center, text_color=color_diff
            ),
            _paragraph_cell(trend, body_style_center),
        ]

    rows = [
        [
            "Metric",
            "Contract A",
            "Contract B",
            "Difference",
            "Trend",
        ],
        _row("overall_risk_score", "Overall Risk"),
        _row("gdpr_readiness_score", "GDPR Readiness"),
        _row("eu_ai_act_readiness_score", "EU AI Act Readiness"),
    ]

    return _make_table(
        rows,
        colWidths=[label_width, value_width, value_width, value_width, value_width],
        styles=styles,
        align="CENTER",
        first_col_bold=False,
    )


def generate_comparison_pdf(
    contract_a_filename: str,
    contract_b_filename: str,
    contract_a_analysis: dict[str, Any],
    contract_b_analysis: dict[str, Any],
    comparison: dict[str, Any],
) -> BytesIO:
    """Generate a professional side-by-side comparison PDF report."""

    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40,
    )

    styles = getSampleStyleSheet()
    usable_width = document.width
    label_width = 1.5 * inch
    two_col_widths = [
        label_width,
        (usable_width - label_width) / 2,
        (usable_width - label_width) / 2,
    ]
    story: list[Any] = []

    findings_a = contract_a_analysis.get("findings", [])
    findings_b = contract_b_analysis.get("findings", [])
    risk_a = contract_a_analysis.get("risk_scores", {})
    risk_b = contract_b_analysis.get("risk_scores", {})
    gdpr_a = contract_a_analysis.get("gdpr_check", {})
    gdpr_b = contract_b_analysis.get("gdpr_check", {})
    ai_a = contract_a_analysis.get("ai_act_check", {})
    ai_b = contract_b_analysis.get("ai_act_check", {})

    score_comparison = comparison.get("score_comparison", {})
    clause_comparison = comparison.get("clause_comparison", {})
    summary = comparison.get("summary", {})

    # Title
    _pdf_heading(story, styles, "Contract Comparison Report", "Heading1")
    _pdf_heading(
        story,
        styles,
        f"{contract_a_filename} <b>vs</b> {contract_b_filename}",
        "Heading2",
    )
    _pdf_body(story, styles, "Generated locally by Legal Clause Analyzer.")
    story.append(Spacer(1, 0.2 * inch))

    # Executive Summary
    _pdf_heading(story, styles, "Executive Summary")
    executive_bullets = [
        (
            f"Contract A ({contract_a_filename}) contains "
            f"<b>{len(findings_a)}</b> detected clause(s)."
        ),
        (
            f"Contract B ({contract_b_filename}) contains "
            f"<b>{len(findings_b)}</b> detected clause(s)."
        ),
        (
            f"Common clause types: <b>{summary.get('common_clauses_count', 0)}</b>; "
            f"added in Contract B: <b>{summary.get('added_clauses_count', 0)}</b>; "
            f"removed from Contract A: <b>{summary.get('removed_clauses_count', 0)}</b>."
        ),
        (
            f"Overall risk: Contract A <b>{risk_a.get('overall_risk_score', 0)}</b>, "
            f"Contract B <b>{risk_b.get('overall_risk_score', 0)}</b> "
            f"(difference {score_comparison.get('overall_risk_score', {}).get('difference', 0)})."
        ),
        (
            f"GDPR readiness: Contract A <b>{risk_a.get('gdpr_readiness_score', 0)}</b>, "
            f"Contract B <b>{risk_b.get('gdpr_readiness_score', 0)}</b>."
        ),
        (
            f"EU AI Act readiness: Contract A <b>{risk_a.get('eu_ai_act_readiness_score', 0)}</b>, "
            f"Contract B <b>{risk_b.get('eu_ai_act_readiness_score', 0)}</b>."
        ),
    ]
    _pdf_bullets(story, styles, executive_bullets)
    _pdf_body(
        story,
        styles,
        (
            "This comparison is intended only as a compliance-readiness "
            "assessment and must not be considered legal advice."
        ),
    )
    story.append(Spacer(1, 0.2 * inch))

    # Contract A Summary
    _pdf_heading(story, styles, "Contract A Summary")
    story.append(
        _contract_summary_table(
            contract_a_filename,
            findings_a,
            risk_a,
            gdpr_a,
            ai_a,
            styles,
            usable_width,
        )
    )
    story.append(Spacer(1, 0.2 * inch))

    # Contract B Summary
    _pdf_heading(story, styles, "Contract B Summary")
    story.append(
        _contract_summary_table(
            contract_b_filename,
            findings_b,
            risk_b,
            gdpr_b,
            ai_b,
            styles,
            usable_width,
        )
    )
    story.append(Spacer(1, 0.2 * inch))

    # Clause Comparison
    _pdf_heading(story, styles, "Clause Comparison")
    clause_rows = [
        ["Item", "Contract A", "Contract B"],
        [
            "Total detected clauses",
            str(len(findings_a)),
            str(len(findings_b)),
        ],
        [
            "Clause types",
            _join_items(clause_comparison.get("contract_a_clause_types", [])),
            _join_items(clause_comparison.get("contract_b_clause_types", [])),
        ],
        [
            "Common clause types",
            _join_items(clause_comparison.get("common", [])),
            _join_items(clause_comparison.get("common", [])),
        ],
        [
            "Added in Contract B",
            "—",
            _join_items(clause_comparison.get("added", [])),
        ],
        [
            "Removed from Contract A",
            _join_items(clause_comparison.get("removed", [])),
            "—",
        ],
    ]
    story.append(
        _make_table(
            clause_rows,
            colWidths=two_col_widths,
            styles=styles,
        )
    )
    story.append(Spacer(1, 0.2 * inch))

    # Risk Score Comparison
    _pdf_heading(story, styles, "Risk Score Comparison")
    story.append(
        _score_comparison_table(
            score_comparison,
            styles,
            usable_width,
        )
    )
    story.append(Spacer(1, 0.2 * inch))

    # GDPR Comparison
    _pdf_heading(story, styles, "GDPR Comparison")
    gdpr_rows = [
        ["Aspect", "Contract A", "Contract B"],
        [
            "Personal data detected",
            "Yes" if gdpr_a.get("personal_data_detected") else "No",
            "Yes" if gdpr_b.get("personal_data_detected") else "No",
        ],
        [
            "Sensitive data detected",
            "Yes" if gdpr_a.get("sensitive_data_detected") else "No",
            "Yes" if gdpr_b.get("sensitive_data_detected") else "No",
        ],
        [
            "Matched personal-data terms",
            _join_items(gdpr_a.get("matched_personal_data_terms", [])),
            _join_items(gdpr_b.get("matched_personal_data_terms", [])),
        ],
        [
            "Matched sensitive-data terms",
            _join_items(gdpr_a.get("matched_sensitive_data_terms", [])),
            _join_items(gdpr_b.get("matched_sensitive_data_terms", [])),
        ],
        [
            "Missing controls count",
            str(len(gdpr_a.get("missing_controls", []))),
            str(len(gdpr_b.get("missing_controls", []))),
        ],
        [
            "Missing controls",
            _join_items(gdpr_a.get("missing_controls", [])),
            _join_items(gdpr_b.get("missing_controls", [])),
        ],
    ]
    story.append(
        _make_table(
            gdpr_rows,
            colWidths=two_col_widths,
            styles=styles,
        )
    )
    _pdf_labeled_bullets(
        story,
        styles,
        "GDPR Issues",
        [
            ("A", gdpr_a.get("issues", [])),
            ("B", gdpr_b.get("issues", [])),
        ],
        fallback="• No GDPR issues detected.",
    )
    _pdf_labeled_bullets(
        story,
        styles,
        "GDPR Recommendations",
        [
            ("A", gdpr_a.get("recommendations", [])),
            ("B", gdpr_b.get("recommendations", [])),
        ],
        fallback="• No GDPR recommendations.",
    )
    story.append(Spacer(1, 0.2 * inch))

    # EU AI Act Comparison
    _pdf_heading(story, styles, "EU AI Act Comparison")
    ai_rows = [
        ["Aspect", "Contract A", "Contract B"],
        [
            "AI Act triggered",
            "Yes" if ai_a.get("ai_act_triggered") else "No",
            "Yes" if ai_b.get("ai_act_triggered") else "No",
        ],
        [
            "Matched AI terms",
            _join_items(ai_a.get("matched_ai_terms", [])),
            _join_items(ai_b.get("matched_ai_terms", [])),
        ],
        [
            "High-risk AI terms",
            _join_items(ai_a.get("matched_high_risk_terms", [])),
            _join_items(ai_b.get("matched_high_risk_terms", [])),
        ],
        [
            "Missing controls count",
            str(len(ai_a.get("missing_controls", []))),
            str(len(ai_b.get("missing_controls", []))),
        ],
        [
            "Missing controls",
            _join_items(ai_a.get("missing_controls", [])),
            _join_items(ai_b.get("missing_controls", [])),
        ],
    ]
    story.append(
        _make_table(
            ai_rows,
            colWidths=two_col_widths,
            styles=styles,
        )
    )
    _pdf_labeled_bullets(
        story,
        styles,
        "EU AI Act Issues",
        [
            ("A", ai_a.get("issues", [])),
            ("B", ai_b.get("issues", [])),
        ],
        fallback="• No EU AI Act issues detected.",
    )
    _pdf_labeled_bullets(
        story,
        styles,
        "EU AI Act Recommendations",
        [
            ("A", ai_a.get("recommendations", [])),
            ("B", ai_b.get("recommendations", [])),
        ],
        fallback="• No EU AI Act recommendations.",
    )
    story.append(Spacer(1, 0.2 * inch))

    _pdf_body(
        story,
        styles,
        (
            "Disclaimer: This report is generated for "
            "compliance-readiness and demonstration purposes only and "
            "does not constitute legal advice."
        ),
    )

    document.build(story)
    buffer.seek(0)
    return buffer


@app.post("/compare-contracts", response_model=CompareContractsResponse)
async def compare_contracts(
    file_a: UploadFile = File(...),
    file_b: UploadFile = File(...),
    use_llm: bool = False,
) -> CompareContractsResponse:
    """Compare two contracts.

    Accepts two PDF or DOCX files, extracts text from each, runs the
    existing full analysis on both, and returns a structured comparison
    of detected clauses and risk/readiness scores.
    """

    global latest_comparison

    bytes_a = await file_a.read()
    bytes_b = await file_b.read()

    try:
        text_a = extract_text(bytes_a, file_a.filename)
    except ValueError as exc:
        raise HTTPException(
            status_code=400,
            detail=f"Contract A: {exc}",
        )

    try:
        text_b = extract_text(bytes_b, file_b.filename)
    except ValueError as exc:
        raise HTTPException(
            status_code=400,
            detail=f"Contract B: {exc}",
        )

    result_a = run_full_analysis(text_a, use_llm)
    result_b = run_full_analysis(text_b, use_llm)

    comparison = compare_analysis_results(result_a, result_b)

    latest_comparison = {
        "file_a": file_a.filename,
        "file_b": file_b.filename,
        "result_a": result_a,
        "result_b": result_b,
        "comparison": comparison,
    }

    return {
        "project": "Legal Clause Analyzer",
        "contract_a_filename": file_a.filename,
        "contract_b_filename": file_b.filename,
        "characters_extracted": {
            "contract_a": len(text_a),
            "contract_b": len(text_b),
            "total": len(text_a) + len(text_b),
        },
        "contract_a_analysis": result_a,
        "contract_b_analysis": result_b,
        "comparison": comparison,
        "disclaimer": (
            "This output is for compliance-readiness and "
            "demonstration only. It is not legal advice."
        ),
    }


@app.get(
    "/download-comparison-report",
    response_class=FileResponse,
    responses={
        200: {
            "description": "Contract comparison PDF report",
            "content": {
                "application/pdf": {
                    "schema": {"type": "string", "format": "binary"},
                }
            },
        },
        404: {"description": "No comparison has been generated yet"},
    },
)
def download_comparison_report(
    background_tasks: BackgroundTasks,
) -> FileResponse:
    """Download the most recent contract comparison as a PDF."""

    global latest_comparison

    if not latest_comparison:
        raise HTTPException(
            status_code=404,
            detail="No comparison available. Please compare two contracts first.",
        )

    pdf = generate_comparison_pdf(
        contract_a_filename=latest_comparison["file_a"],
        contract_b_filename=latest_comparison["file_b"],
        contract_a_analysis=latest_comparison["result_a"],
        contract_b_analysis=latest_comparison["result_b"],
        comparison=latest_comparison["comparison"],
    )

    tmp_dir = r"C:\Users\Lenovo\AppData\Local\Temp\opencode"
    with tempfile.NamedTemporaryFile(
        delete=False,
        suffix=".pdf",
        dir=tmp_dir,
    ) as tmp:
        tmp.write(pdf.getvalue())
        tmp_path = tmp.name

    background_tasks.add_task(os.remove, tmp_path)

    return FileResponse(
        tmp_path,
        media_type="application/pdf",
        filename="Comparison_Report.pdf",
    )


def generate_pdf_report(
    findings,
    risk_scores,
    ai_act_check,
    gdpr_check,
    llm_summary,
) -> BytesIO:

    buffer = BytesIO()

    document = SimpleDocTemplate(
        buffer,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40,
    )
    
    styles = getSampleStyleSheet()
    story = []

    story.append(
        Paragraph(
            "Legal Clause Analyzer Report",
            styles["Heading1"],
        )
    )

    story.append(
        Paragraph(
            "Professional Compliance Report",
            styles["Heading2"],
        )
    )

    story.append(
        Paragraph(
            "Generated locally by Legal Clause Analyzer.",
            styles["BodyText"],
        )
    )

    story.append(
        Paragraph(
            "<br/><b>Risk Summary</b>",
            styles["Heading2"],
        )
    )

    overall = risk_scores["overall_risk_score"]
    gdpr = risk_scores["gdpr_readiness_score"]
    ai = risk_scores["eu_ai_act_readiness_score"]

    if overall >= 70:
        overall_color = colors.red
    elif overall >= 40:
        overall_color = colors.orange
    else:
        overall_color = colors.green

    if gdpr >= 80:
        gdpr_color = colors.green
    elif gdpr >= 50:
        gdpr_color = colors.orange
    else:
        gdpr_color = colors.red

    if ai >= 80:
        ai_color = colors.green
    elif ai >= 50:
        ai_color = colors.orange
    else:
        ai_color = colors.red

    score_table = [
        ["Metric", "Score"],
        ["Overall Risk", str(overall)],
        ["GDPR Readiness", str(gdpr)],
        ["EU AI Act Readiness", str(ai)],
    ]

    table = Table(
        score_table,
        colWidths=[3.8 * inch, 2 * inch],
    )

    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.darkblue),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("BACKGROUND", (0, 1), (-1, -1), colors.beige),
                ("GRID", (0, 0), (-1, -1), 1, colors.grey),
                ("BOX", (0, 0), (-1, -1), 1, colors.black),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (1, 1), (1, 3), "Helvetica-Bold"),
                ("TEXTCOLOR", (1, 1), (1, 1), overall_color),
                ("TEXTCOLOR", (1, 2), (1, 2), gdpr_color),
                ("TEXTCOLOR", (1, 3), (1, 3), ai_color),
            ]
        )
    )
    
    story.append(table)

    story.append(
        Paragraph(
            "<br/><b>Executive Summary</b>",
            styles["Heading2"],
        )
    )

    story.append(
        Paragraph(
            (
                "This report was generated locally using FastAPI and "
                "Llama 3. No contract text was transmitted to external "
                "cloud services."
            ),
            styles["BodyText"],
        )
    )

    story.append(
        Paragraph(
            (
                "This report is intended only as a compliance-readiness "
                "assessment and must not be considered legal advice."
            ),
            styles["BodyText"],
        )
    )

    story.append(
        Paragraph(
            "<br/><b>Detected Clauses</b>",
            styles["Heading2"],
        )
    )

    if findings:

        for finding in findings:

            story.append(
                Paragraph(
                    (
                        "<b>"
                        + finding["clause_type"]
                        + "</b>"
                        + " — Risk Level: "
                        + finding["risk_level"]
                    ),
                    styles["BodyText"],
                )
            )

            if "explanation" in finding:

                story.append(
                    Paragraph(
                        finding["explanation"],
                        styles["BodyText"],
                    )
                )

    else:

        story.append(
            Paragraph(
                "No legal clauses were detected.",
                styles["BodyText"],
            )
        )

    story.append(
        Paragraph(
            "<br/><b>GDPR Findings</b>",
            styles["Heading2"],
        )
    )

    for issue in gdpr_check["issues"]:

        story.append(
            Paragraph(
                "• " + issue,
                styles["BodyText"],
            )
        )

    story.append(
        Paragraph(
            "<br/><b>GDPR Recommendations</b>",
            styles["Heading3"],
        )
    )

    for recommendation in gdpr_check["recommendations"]:

        story.append(
            Paragraph(
                "• " + recommendation,
                styles["BodyText"],
            )
        )

    story.append(
        Paragraph(
            "<br/><b>EU AI Act Findings</b>",
            styles["Heading2"],
        )
    )

    for issue in ai_act_check["issues"]:

        story.append(
            Paragraph(
                "• " + issue,
                styles["BodyText"],
            )
        )

    story.append(
        Paragraph(
            "<br/><b>EU AI Act Recommendations</b>",
            styles["Heading3"],
        )
    )

    for recommendation in ai_act_check["recommendations"]:

        story.append(
            Paragraph(
                "• " + recommendation,
                styles["BodyText"],
            )
        )
    
    if llm_summary:
        story.append(
            Paragraph(
                "<br/><b>LLM Summary</b>",
                styles["Heading2"],
            )
        )

        story.append(
            Paragraph(
                llm_summary.replace("\n", "<br/>"),
                styles["BodyText"],
            )
        )
    
    
    document.build(story)
    buffer.seek(0)
    return buffer


@app.get("/download-report")
def download_report():

    global latest_analysis

    if not latest_analysis:
        return {
        "error": "No analysis available. Please analyze a PDF first."
    }

    findings = latest_analysis["findings"]
    risk_scores = latest_analysis["risk_scores"]
    ai_act_check = latest_analysis["ai_act_check"]
    gdpr_check = latest_analysis["gdpr_check"]
    llm_summary = latest_analysis["llm_summary"]

    pdf = generate_pdf_report(
        findings,
        risk_scores,
        ai_act_check,
        gdpr_check,
        llm_summary,
    )

    return StreamingResponse(
        pdf,
        media_type="application/pdf",
        headers={
            "Content-Disposition":
            "attachment; filename=Legal_Report.pdf"
        },
    )

      
@app.get("/")
def read_root() -> dict[str, str]:
    return {
        "status": "ok",
        "project": "Legal Clause Analyzer",
        "docs": "/docs",
    }


@app.post("/analyze")
def analyze_contract(request: AnalyzeRequest) -> dict[str, Any]:
    result = run_full_analysis(request.contract_text, request.use_llm)

    return {
        "project": "Legal Clause Analyzer",
        "privacy_note": (
            "This prototype does not store contract text. "
            "Use synthetic or anonymized data for demos."
        ),
        "clause_findings": result["findings"],
        "ai_act_compliance_check": result["ai_act_check"],
        "gdpr_privacy_check": result["gdpr_check"],
        "risk_scores": result["risk_scores"],
        "llm_summary": result["llm_summary"],
        "disclaimer": (
            "This output is for compliance-readiness and product "
            "demonstration only. It is not legal advice."
        ),
    } 